"""Render the confidence dict to a human-readable Word (.docx) report.

The flags table carries a 4th "What to check" column from reviewer_guide — this guidance lives in the
Word doc ONLY and is never written back into the confidence JSON (the JSON stays the machine-readable
source of truth). Renderer is pure: it reads the confidence dict, mutates nothing. python-docx only.
"""
from __future__ import annotations
import io

from docx import Document

from . import reviewer_guide


def _bold_cell(cell, text: str) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True


def _counts(flags):
    hard = sum(1 for f in flags if str(f.get("level", "")).upper() == "HARD")
    soft = sum(1 for f in flags if str(f.get("level", "")).upper() == "SOFT")
    return hard, soft


def render_confidence_docx(confidence: dict) -> bytes:
    """Confidence dict -> .docx bytes. Header summary + the flags table
    (Level | Item | Message | What to check). The racking family collapses to one row and any
    unmapped flag shows the visible reviewer-guidance fallback (both via reviewer_guide)."""
    c = confidence or {}
    flags = c.get("FLAGS_FOR_HUMAN_REVIEW", []) or []
    proj = c.get("project", {}) or {}
    planset = c.get("planset", {}) or {}
    hard, soft = _counts(flags)

    doc = Document()
    doc.add_heading("BOM Confidence Report", level=0)

    title = str(proj.get("name") or "").strip()
    if proj.get("number"):
        title = f"{title}  (#{proj.get('number')})".strip()
    if title:
        doc.add_paragraph(title)

    summary = doc.add_paragraph()
    summary.add_run("Mode: ").bold = True
    summary.add_run(str(c.get("mode") or "—"))
    if planset.get("file_name"):
        p = doc.add_paragraph()
        p.add_run("Planset: ").bold = True
        p.add_run(f"{planset.get('file_name')}  (rev {planset.get('revision', '—')})")
    fl = doc.add_paragraph()
    fl.add_run("Flags: ").bold = True
    fl.add_run(f"{hard} hard / {soft} soft")
    if hard:
        fl.add_run("  — HARD FLAGS PRESENT; review required before use.").bold = True

    doc.add_heading("Flags for human review", level=1)
    if not flags:
        doc.add_paragraph("No flags — clean run.")
        return _to_bytes(doc)

    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = "Light Grid Accent 1"
    except Exception:  # noqa: BLE001 — style is cosmetic; default table is fine if the theme lacks it
        pass
    hdr = table.rows[0].cells
    for i, h in enumerate(("Level", "Item", "Message", "What to check")):
        _bold_cell(hdr[i], h)

    for level, item, message, location, check in reviewer_guide.reviewer_rows(flags):
        cells = table.add_row().cells
        cells[0].text = str(level)
        cells[1].text = str(item)
        cells[2].text = str(message)
        # "What to check" = location + guidance; the fallback marker has no real location.
        cells[3].text = (f"{location} — {check}"
                         if location and location != reviewer_guide.FALLBACK_LOCATION else str(check))
    return _to_bytes(doc)


def _to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
